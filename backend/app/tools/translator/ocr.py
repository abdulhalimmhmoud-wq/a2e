"""استخراج النص من المستندات الممسوحة ضوئيًا عبر قدرة الرؤية في Claude.

ليه مش Tesseract؟ لأنه محتاج تثبيت منفصل وحزم لغة عربية، ودقّته على
العربي المكتوب بخطوط متنوعة أو الممسوح بجودة متوسطة أقل بوضوح. النموذج
بيقرا الصفحة زي ما الإنسان بيقراها، وبيفهم البنية (عناوين/فقرات/جداول)
مش بس الحروف.

المخرَج ملف Word مبني من الصفر بالنص المستخرج — بعد كده كل الخط
بيشتغل عليه زي أي ملف Word عادي، فمفيش مسار منفصل نصونه.

التكلفة: صفحة A4 بدقة 150 نقطة/بوصة ≈ 2500–3000 توكن دخل. يعني حوالي
نصف سنت للصفحة على Sonnet 5 — بتظهر في حاسبة التكلفة زي أي نداء تاني.
"""
from __future__ import annotations

import base64
import json
import logging
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.config import settings
from app.tools.translator.costing import compute_cost
from app.tools.translator.engine import Usage

logger = logging.getLogger(__name__)

# دقّة المسح. أعلى = أدق وأغلى. 150 توازن عملي للمستندات المكتوبة.
DEFAULT_DPI = 150

# الواجهة البرمجية بترفض أي صورة أكبر من 10 ميجا بعد ترميز base64.
# بنشتغل تحت الحد بهامش واسع لأن حجم الترميز بيختلف من صفحة للتانية
# حسب كثافة الحبر والضوضاء في المسح.
MAX_IMAGE_B64 = 4_500_000

# النموذج بيصغّر أي صورة ضلعها الأطول أكبر من ~1568 بكسل قبل ما يقراها،
# فالرسم فوق كده بيبعت بيانات على الشبكة وتترمي. بنسيب هامش فوق الحد
# عشان التصغير عنده يلاقي تفاصيل يشتغل عليها.
MAX_EDGE_PX = 2200

# الصفحة الممسوحة صورة فوتوغرافية فيها ضوضاء، و PNG بلا فقد أسوأ صيغة
# ليها: نفس الصفحة طلعت 25.9 ميجا بـ PNG مقابل 3.6 ميجا بـ JPEG بنفس
# الدقة بالظبط. الفرق ده هو اللي كان بيكسّر الحد.
JPEG_QUALITY = 85
_MEDIA_TYPE = "image/jpeg"

_SYSTEM = """\
You transcribe scanned document pages into structured text.

Rules:
1. Transcribe **exactly** what is on the page, in the original language.
   Do not translate, summarise, correct, or complete anything.
2. Preserve reading order as a human reader of that language would: top to
   bottom, and right-to-left for Arabic, Hebrew, Persian, and Urdu pages.
   Multi-column layouts are read column by column.
3. Reproduce every number, date, code, and identifier exactly as printed.
4. Classify each block: `heading` for titles and section headers,
   `list_item` for bulleted or numbered entries, `table_row` for a row of a
   table (put each cell in `cells`), `paragraph` for everything else.
5. For a `table_row`, leave `text` empty and fill `cells`. For every other
   kind, fill `text` and leave `cells` empty.
6. Skip page furniture that carries no content: page numbers alone,
   decorative rules, scan artefacts.
7. If a word is genuinely illegible, write `[…]` in its place rather than
   guessing. A guessed word in a medical or legal document is worse than a
   visible gap.
"""

_SCHEMA = {
    "type": "object",
    "properties": {
        "blocks": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "kind": {
                        "type": "string",
                        "enum": ["heading", "paragraph", "list_item", "table_row"],
                    },
                    "text": {"type": "string"},
                    "cells": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["kind", "text", "cells"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["blocks"],
    "additionalProperties": False,
}


@dataclass
class Block:
    kind: str
    text: str = ""
    cells: list[str] = field(default_factory=list)
    page: int = 0


@dataclass
class OcrResult:
    blocks: list[Block] = field(default_factory=list)
    usage: Usage = field(default_factory=Usage)
    pages: int = 0
    failed_pages: list[int] = field(default_factory=list)


def _b64_size(raw: int) -> int:
    """حجم البيانات بعد ترميز base64 — الحد بيتحسب على المرمَّز."""
    return (raw + 2) // 3 * 4


def render_page(page, dpi: int = DEFAULT_DPI, budget: int = MAX_IMAGE_B64) -> bytes:
    """رسم صفحة واحدة كصورة JPEG مضمونة إنها تحت حد الحجم.

    مقاسات الصفحات بتتفاوت بشكل كبير: الصفحة اللي كشفت العيب ده كانت
    ٢٥.٨×٣٦.٨ بوصة (مقاس مخططات)، فدقة ثابتة عليها بتطلع ٣٨٦٩×٥٥٢٣
    بكسل. أي دقة ثابتة هتنفجر على مقاس ما، فالقياس هنا على الحجم
    الفعلي بعد الترميز مش على الدقة.
    """
    zoom = dpi / 72.0
    longest = max(page.rect.width, page.rect.height) * zoom
    if longest > MAX_EDGE_PX:
        zoom *= MAX_EDGE_PX / longest

    # بننزّل الجودة الأول قبل ما ننزّل الدقة: خفض الجودة بيوفّر حجمًا
    # كتير من غير ما يضيّع حروفًا، أما خفض الدقة فبيصغّر الحروف نفسها
    # وده اللي بيأذي القراءة.
    data = b""
    for _ in range(8):
        pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        for quality in (JPEG_QUALITY, 70, 55):
            data = pixmap.tobytes("jpeg", jpg_quality=quality)
            if _b64_size(len(data)) <= budget:
                return data
        zoom *= 0.7

    logger.warning(
        "صفحة فضلت %.1f ميجا بعد التصغير المتكرر — هتتبعت زي ما هي",
        _b64_size(len(data)) / 1e6,
    )
    return data


def render_pages(path: Path, dpi: int = DEFAULT_DPI) -> list[bytes]:
    """تحويل صفحات الـ PDF لصور JPEG تحت حد حجم الواجهة البرمجية."""
    images: list[bytes] = []
    with fitz.open(str(path)) as document:
        for page in document:
            images.append(render_page(page, dpi))
    return images


def _read_page(client, model: str, image: bytes, page_number: int) -> tuple[list[Block], Usage]:
    response = client.messages.create(
        model=model,
        max_tokens=8000,
        system=[{"type": "text", "text": _SYSTEM, "cache_control": {"type": "ephemeral"}}],
        thinking={"type": "adaptive"},
        output_config={
            "effort": "medium",
            "format": {"type": "json_schema", "schema": _SCHEMA},
        },
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": _MEDIA_TYPE,
                            "data": base64.standard_b64encode(image).decode(),
                        },
                    },
                    {
                        "type": "text",
                        "text": f"Transcribe page {page_number} of this scanned document.",
                    },
                ],
            }
        ],
    )

    raw = response.usage
    usage = Usage(
        input_tokens=getattr(raw, "input_tokens", 0) or 0,
        output_tokens=getattr(raw, "output_tokens", 0) or 0,
        cache_read_tokens=getattr(raw, "cache_read_input_tokens", 0) or 0,
        cache_write_tokens=getattr(raw, "cache_creation_input_tokens", 0) or 0,
        calls=1,
    )
    usage.cost_usd = compute_cost(
        model,
        input_tokens=usage.input_tokens,
        output_tokens=usage.output_tokens,
        cache_read_tokens=usage.cache_read_tokens,
        cache_write_tokens=usage.cache_write_tokens,
    )

    if response.stop_reason == "refusal":
        return [], usage

    text = next((b.text for b in response.content if b.type == "text"), "")
    if not text.strip():
        return [], usage

    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        logger.warning("رد OCR غير قابل للقراءة في صفحة %d", page_number)
        return [], usage

    blocks = [
        Block(
            kind=item.get("kind", "paragraph"),
            text=item.get("text", ""),
            cells=item.get("cells", []) or [],
            page=page_number,
        )
        for item in parsed.get("blocks", [])
    ]
    return blocks, usage


def read_document(
    path: Path,
    api_key: str | None = None,
    model: str | None = None,
    dpi: int = DEFAULT_DPI,
    concurrency: int | None = None,
    progress=None,
) -> OcrResult:
    """قراءة كل صفحات المستند الممسوح.

    الصفحات مستقلة تمامًا فبتتقرا بالتوازي. أول صفحة لوحدها عشان تكتب
    التعليمات في الكاش وباقي الصفحات تقراها بعُشر السعر.
    """
    import anthropic

    key = api_key or settings.anthropic_api_key
    if not key:
        raise RuntimeError("الـ OCR محتاج مفتاح Anthropic")

    client = anthropic.Anthropic(api_key=key)
    model = model or settings.default_model
    concurrency = concurrency or settings.translation_concurrency

    images = render_pages(path, dpi)
    result = OcrResult(pages=len(images))
    if not images:
        return result

    collected: dict[int, list[Block]] = {}

    def read(index: int) -> tuple[int, list[Block], Usage]:
        blocks, usage = _read_page(client, model, images[index], index + 1)
        return index, blocks, usage

    def absorb(item: tuple[int, list[Block], Usage]) -> None:
        index, blocks, usage = item
        collected[index] = blocks
        result.usage.add(usage)
        if not blocks:
            result.failed_pages.append(index + 1)
        if progress:
            progress(len(collected), len(images))

    # الصفحة الأولى لوحدها: بتكتب الكاش وباقي الصفحات بتقرا منه
    absorb(read(0))

    if len(images) > 1:
        workers = max(1, min(concurrency, len(images) - 1))
        with ThreadPoolExecutor(max_workers=workers, thread_name_prefix="ocr") as pool:
            for item in pool.map(read, range(1, len(images))):
                absorb(item)

    for index in sorted(collected):
        result.blocks.extend(collected[index])

    return result


# ---------------------------------------------------------------------------
# بناء ملف Word من نتيجة القراءة
# ---------------------------------------------------------------------------
def _set_direction(paragraph, rtl: bool) -> None:
    """ضبط اتجاه الفقرة حسب لغة **المصدر**.

    الملف المبني هنا هو مدخل الترجمة، فبيتبني باتجاه المصدر. قلب
    الاتجاه للغة الهدف بيحصل وقت التصدير زي أي ملف تاني.
    """
    if rtl:
        ppr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
        ppr.append(OxmlElement("w:bidi"))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_docx(
    result: OcrResult,
    output: Path,
    title: str = "",
    source_lang: str = "ar",
) -> tuple[Path, list[dict]]:
    """تحويل نتيجة القراءة لملف Word جاهز للترجمة.

    بيرجّع كمان **خريطة الصفحات**: لكل وحدة نص بتتكتب في الملف،
    رقم صفحتها في الأصل ونصّها. الخريطة دي هي اللي بتخلي التصدير
    يقدر يرجّع الترجمة لمكانها على الصفحة الأصلية بدل ما يبني ملفًا
    جديدًا بيضيّع الشعارات والزخارف.
    """
    from app.tools.translator.langs import is_rtl

    rtl = is_rtl(source_lang)
    document = Document()

    def _set_rtl(paragraph) -> None:
        _set_direction(paragraph, rtl)

    if title:
        heading = document.add_heading(title, level=0)
        _set_rtl(heading)

    # خريطة الصفحات: عنصر لكل وحدة نص **بالترتيب اللي بتتكتب بيه**،
    # عشان تطابق ترتيب الوحدات اللي الاستخراج هيقراها من نفس الملف.
    layout: list[dict] = []
    pending_rows: list[tuple[list[str], int]] = []

    def flush_table() -> None:
        if not pending_rows:
            return
        width = max(len(row) for row, _ in pending_rows)
        table = document.add_table(rows=len(pending_rows), cols=width)
        table.style = "Table Grid"
        for row_index, (row, page) in enumerate(pending_rows):
            for col_index in range(width):
                cell = table.cell(row_index, col_index)
                cell.text = row[col_index] if col_index < len(row) else ""
                for paragraph in cell.paragraphs:
                    _set_rtl(paragraph)
                layout.append({"page": page, "text": cell.text, "kind": "cell"})
        pending_rows.clear()

    current_page = 0
    for block in result.blocks:
        if block.page != current_page:
            flush_table()
            current_page = block.page

        if block.kind == "table_row":
            if block.cells:
                pending_rows.append((block.cells, block.page))
            continue

        flush_table()
        text = block.text.strip()
        if not text:
            continue

        if block.kind == "heading":
            paragraph = document.add_heading(text, level=2)
        elif block.kind == "list_item":
            paragraph = document.add_paragraph(text, style="List Bullet")
        else:
            paragraph = document.add_paragraph(text)
        _set_rtl(paragraph)
        layout.append({"page": block.page, "text": text, "kind": block.kind})

    flush_table()

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output, layout
