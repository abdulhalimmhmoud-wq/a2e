"""معالج Word (DOCX) — استخراج ودمج.

الاستخراج والدمج في نفس الموديول عن قصد: الاتنين لازم يمشوا في المستند
بنفس الترتيب بالظبط عشان الـ unit_key يرجّعنا لنفس الفقرة. أي اختلاف
بين المشيتين = نص بيتحط في المكان الغلط.

اللي بنغطّيه:
  المتن · الجداول (والمتداخلة) · الرؤوس والتذييلات · الحواشي السفلية
  والختامية · التعليقات · مربعات النص · الروابط · القوائم المرقّمة

تعامُلنا مع التنسيق:
  الجملة الواحدة في Word ممكن تتقسّم على عدة "runs" بسبب التنسيق.
  بندمج الـ runs المتشابهة في مجموعات، وبنمثّل المجموعات المختلفة عن
  التنسيق الغالب بوسوم <g1>...</g1> يحافظ عليها النموذج أثناء الترجمة.
  لو النموذج ضيّع وسم، بنرجع لوضع آمن (كل النص بالتنسيق الغالب) —
  التنسيق بيتبسّط لكن **النص عمره ما بيضيع**.
"""
from __future__ import annotations

import copy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from app.tools.translator.formats.base import (
    ExtractionResult,
    TextUnit,
    parse_tagged_text,
    tags_in,
)

# أنواع المحتوى اللي فيها نص قابل للترجمة
_TEXT_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml": "header",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml": "footer",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml": "footnote",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml": "endnote",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml": "comment",
}


# ---------------------------------------------------------------------------
# الوصول للأجزاء (parts) داخل ملف Word
# ---------------------------------------------------------------------------
@dataclass
class _PartHandle:
    part_id: str
    kind: str
    root: etree._Element
    _part: object
    _is_xml_part: bool

    def flush(self) -> None:
        """كتابة التعديلات لجزء غير مُدار من python-docx."""
        if not self._is_xml_part:
            self._part._blob = etree.tostring(  # noqa: SLF001
                self.root, xml_declaration=True, encoding="UTF-8", standalone=True
            )


def _iter_parts(document) -> list[_PartHandle]:
    """كل أجزاء المستند اللي فيها نص، بترتيب حتمي ثابت.

    المتن أولًا، وبعدين باقي الأجزاء مرتّبة أبجديًا باسم الجزء —
    الترتيب الأبجدي مهم عشان يفضل ثابت بين الاستخراج والدمج.
    """
    handles: list[_PartHandle] = [
        _PartHandle("document", "body", document.element, document.part, True)
    ]

    others: list[tuple[str, _PartHandle]] = []
    for part in document.part.package.iter_parts():
        kind = _TEXT_CONTENT_TYPES.get(part.content_type)
        if kind is None:
            continue

        partname = str(part.partname)
        part_id = partname.rsplit("/", 1)[-1].removesuffix(".xml")

        element = getattr(part, "element", None)
        if element is not None:
            handle = _PartHandle(part_id, kind, element, part, True)
        else:
            # جزء غير مُدار (الحواشي/التعليقات) — نفكّه ونعيد كتابته يدويًا
            try:
                root = etree.fromstring(part.blob)
            except Exception:
                continue
            handle = _PartHandle(part_id, kind, root, part, False)

        others.append((partname, handle))

    others.sort(key=lambda item: item[0])
    handles.extend(handle for _, handle in others)
    return handles


# ---------------------------------------------------------------------------
# تجميع الـ runs
# ---------------------------------------------------------------------------
@dataclass
class _RunGroup:
    """مجموعة runs متتالية بنفس التنسيق ونفس الحاوية."""

    sig: str                       # XML الخاص بـ w:rPr (بصمة التنسيق)
    container: etree._Element | None  # عنصر w:hyperlink لو الـ run جوّه رابط
    runs: list[etree._Element] = field(default_factory=list)
    text: str = ""


def _own_runs(paragraph: etree._Element) -> list[etree._Element]:
    """runs الفقرة دي بس — مش اللي جوّه فقرات متداخلة (مربعات النص).

    من غير الفلترة دي، نص مربع النص هيتحسب مرتين: مرة مع الفقرة الأم
    ومرة مع فقرته الخاصة.
    """
    result = []
    for run in paragraph.iter(qn("w:r")):
        ancestor = run.getparent()
        while ancestor is not None and ancestor.tag != qn("w:p"):
            ancestor = ancestor.getparent()
        if ancestor is paragraph:
            result.append(run)
    return result


def _run_text(run: etree._Element) -> str:
    """نص الـ run مع تحويل الفواصل لمحارف عادية."""
    out: list[str] = []
    for child in run:
        tag = child.tag
        if tag == qn("w:t"):
            out.append(child.text or "")
        elif tag == qn("w:tab"):
            out.append("\t")
        elif tag in (qn("w:br"), qn("w:cr")):
            out.append("\n")
        elif tag == qn("w:noBreakHyphen"):
            out.append("-")
    return "".join(out)


def _run_signature(run: etree._Element) -> str:
    """بصمة تنسيق الـ run = محتوى w:rPr كنص XML."""
    rpr = run.find(qn("w:rPr"))
    if rpr is None:
        return ""
    return etree.tostring(rpr, encoding="unicode")


def _hyperlink_container(run: etree._Element) -> etree._Element | None:
    parent = run.getparent()
    if parent is not None and parent.tag == qn("w:hyperlink"):
        return parent
    return None


def _group_runs(paragraph: etree._Element) -> list[_RunGroup]:
    """دمج الـ runs المتتالية المتشابهة في مجموعات."""
    groups: list[_RunGroup] = []
    for run in _own_runs(paragraph):
        sig = _run_signature(run)
        container = _hyperlink_container(run)
        text = _run_text(run)

        if groups and groups[-1].sig == sig and groups[-1].container is container:
            groups[-1].runs.append(run)
            groups[-1].text += text
        else:
            groups.append(_RunGroup(sig=sig, container=container, runs=[run], text=text))
    return groups


def _dominant_index(groups: list[_RunGroup]) -> int | None:
    """المجموعة الغالبة = أطول نص بين المجموعات اللي مش داخل رابط.

    مجموعات الروابط لازم تفضل موسومة دايمًا، عشان نعرف نحط النص
    المترجم جوّه عنصر الرابط الصحيح وقت الدمج.
    """
    best_index, best_len = None, -1
    for index, group in enumerate(groups):
        if group.container is not None:
            continue
        if len(group.text) > best_len:
            best_index, best_len = index, len(group.text)
    return best_index


def _build_tagged_text(groups: list[_RunGroup]) -> tuple[str, dict[str, str]]:
    """تحويل المجموعات لنص موسوم + خريطة الوسوم."""
    dominant = _dominant_index(groups)
    parts: list[str] = []
    placeholders: dict[str, str] = {}
    tag_counter = 0

    for index, group in enumerate(groups):
        if not group.text:
            continue
        if index == dominant:
            parts.append(group.text)
        else:
            tag_counter += 1
            tag = str(tag_counter)
            placeholders[tag] = group.sig
            parts.append(f"<g{tag}>{group.text}</g{tag}>")

    return "".join(parts), placeholders


def _tag_to_group_index(groups: list[_RunGroup]) -> dict[str, int]:
    """إعادة بناء خريطة (رقم الوسم → المجموعة) بنفس قاعدة الاستخراج.

    لازم تطابق ترتيب `_build_tagged_text` بالضبط.
    """
    dominant = _dominant_index(groups)
    mapping: dict[str, int] = {}
    tag_counter = 0
    for index, group in enumerate(groups):
        if not group.text:
            continue
        if index == dominant:
            continue
        tag_counter += 1
        mapping[str(tag_counter)] = index
    return mapping


# ---------------------------------------------------------------------------
# وصف الموضع للمراجع
# ---------------------------------------------------------------------------
def _describe(paragraph: etree._Element, part_kind: str, ordinal: int) -> str:
    """وصف مقروء لمكان الفقرة يظهر في شاشة المراجعة."""
    labels = {
        "header": "ترويسة",
        "footer": "تذييل",
        "footnote": "حاشية سفلية",
        "endnote": "حاشية ختامية",
        "comment": "تعليق",
    }
    if part_kind in labels:
        return labels[part_kind]

    # هل الفقرة جوّه مربع نص؟
    node = paragraph.getparent()
    in_textbox = False
    cell = row = table = None
    while node is not None:
        if node.tag == qn("w:txbxContent"):
            in_textbox = True
        elif node.tag == qn("w:tc") and cell is None:
            cell = node
        elif node.tag == qn("w:tr") and row is None:
            row = node
        elif node.tag == qn("w:tbl") and table is None:
            table = node
        node = node.getparent()

    if in_textbox:
        return "مربع نص"

    if table is not None and row is not None and cell is not None:
        try:
            row_index = list(table.iter(qn("w:tr"))).index(row) + 1
            col_index = list(row.iter(qn("w:tc"))).index(cell) + 1
            return f"جدول — صف {row_index} عمود {col_index}"
        except ValueError:
            return "جدول"

    return f"فقرة {ordinal}"


def _paragraph_kind(paragraph: etree._Element, part_kind: str) -> str:
    if part_kind in ("header", "footer", "footnote", "endnote", "comment"):
        return part_kind
    node = paragraph.getparent()
    while node is not None:
        if node.tag == qn("w:tc"):
            return "cell"
        if node.tag == qn("w:txbxContent"):
            return "textbox"
        node = node.getparent()
    return "paragraph"


# ---------------------------------------------------------------------------
# الاستخراج
# ---------------------------------------------------------------------------
def extract(path: Path) -> ExtractionResult:
    document = Document(str(path))
    units: list[TextUnit] = []
    order = 0

    for handle in _iter_parts(document):
        paragraphs = list(handle.root.iter(qn("w:p")))
        for index, paragraph in enumerate(paragraphs):
            groups = _group_runs(paragraph)
            text, placeholders = _build_tagged_text(groups)
            if not text.strip():
                continue

            units.append(
                TextUnit(
                    unit_key=f"docx:{handle.part_id}:p:{index:05d}",
                    text=text,
                    kind=_paragraph_kind(paragraph, handle.kind),
                    location=_describe(paragraph, handle.kind, index + 1),
                    order_index=order,
                    placeholders=placeholders,
                )
            )
            order += 1

    return ExtractionResult(
        units=units,
        page_count=_estimate_pages(document),
        meta={"parts": len(_iter_parts(document))},
    )


def _estimate_pages(document) -> int:
    """عدد الصفحات التقريبي.

    Word مابيخزّنش عدد الصفحات الحقيقي في الملف — بيتحسب وقت العرض.
    القيمة المخزّنة في app.xml بتكون من آخر مرة اتحفظ فيها الملف ببرنامج
    Word، فبنستخدمها لو موجودة وإلا بنقدّر بالكلمات.
    """
    try:
        core = document.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties"
        )
        root = etree.fromstring(core.blob)
        for tag in root.iter():
            if tag.tag.endswith("}Pages") and tag.text and tag.text.isdigit():
                return int(tag.text)
    except Exception:
        pass

    words = sum(len(p.text.split()) for p in document.paragraphs)
    return max(1, round(words / 250))


# ---------------------------------------------------------------------------
# الدمج — كتابة الترجمة في نسخة من الملف الأصلي
# ---------------------------------------------------------------------------
def _make_run(template: etree._Element, text: str) -> etree._Element:
    """إنشاء run جديد بنفس تنسيق القالب ونص جديد."""
    run = copy.deepcopy(template)
    # نشيل كل المحتوى ونسيب التنسيق بس
    for child in list(run):
        if child.tag != qn("w:rPr"):
            run.remove(child)

    # تقسيم النص على الفواصل وإعادة بنائها كعناصر Word صحيحة
    buffer = ""
    for char in text:
        if char in ("\n", "\t"):
            if buffer:
                _append_text(run, buffer)
                buffer = ""
            run.append(OxmlElement("w:br" if char == "\n" else "w:tab"))
        else:
            buffer += char
    if buffer:
        _append_text(run, buffer)

    return run


def _append_text(run: etree._Element, text: str) -> None:
    node = OxmlElement("w:t")
    # ضروري جدًا: من غيره Word بيبلع المسافات في أول/آخر النص
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)


def _rewrite_paragraph(
    paragraph: etree._Element, groups: list[_RunGroup], target: str
) -> None:
    """استبدال محتوى الفقرة بالنص المترجم مع الحفاظ على التنسيق."""
    if not groups:
        return

    mapping = _tag_to_group_index(groups)
    dominant = _dominant_index(groups)
    pieces = parse_tagged_text(target)

    # فحص السلامة: لو النموذج ضيّع وسمًا، نرجع للوضع الآمن
    expected = set(mapping.keys())
    present = tags_in(target)
    safe_mode = not expected.issubset(present) or dominant is None

    if safe_mode:
        pieces = [(None, "".join(text for _, text in pieces))]

    # نجمّع القطع المتتالية اللي في نفس المجموعة
    assignments: list[tuple[int, str]] = []
    for tag, text in pieces:
        if not text:
            continue
        if tag is not None and tag in mapping:
            group_index = mapping[tag]
        else:
            group_index = dominant if dominant is not None else 0
        if assignments and assignments[-1][0] == group_index:
            assignments[-1] = (group_index, assignments[-1][1] + text)
        else:
            assignments.append((group_index, text))

    # نبني سلسلة العناصر الجديدة بترتيب النص المترجم
    new_children: list[etree._Element] = []
    open_link: etree._Element | None = None
    open_link_source: etree._Element | None = None

    for group_index, text in assignments:
        group = groups[group_index]
        run = _make_run(group.runs[0], text)

        if group.container is not None:
            if open_link_source is not group.container:
                # نسخة جديدة من عنصر الرابط (بتحافظ على r:id وبالتالي على الوجهة)
                open_link = copy.deepcopy(group.container)
                for child in list(open_link):
                    open_link.remove(child)
                open_link_source = group.container
                new_children.append(open_link)
            open_link.append(run)
        else:
            open_link = None
            open_link_source = None
            new_children.append(run)

    # نشيل المحتوى القديم مع الحفاظ على pPr والإشارات المرجعية
    keep_tags = {qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}
    preserved = [child for child in paragraph if child.tag in keep_tags]
    for child in list(paragraph):
        paragraph.remove(child)

    ppr = [c for c in preserved if c.tag == qn("w:pPr")]
    bookmarks = [c for c in preserved if c.tag != qn("w:pPr")]
    for element in ppr:
        paragraph.append(element)
    for element in new_children:
        paragraph.append(element)
    for element in bookmarks:
        paragraph.append(element)


def merge(
    source_path: Path,
    output_path: Path,
    translations: dict[str, str],
    target_rtl: bool = False,
    lang_tag: str = "en-US",
) -> None:
    """كتابة الترجمات في نسخة من الملف الأصلي.

    translations: {unit_key: النص المترجم الموسوم}
    target_rtl:   هل لغة الهدف بتتكتب من اليمين لليسار؟
    """
    document = Document(str(source_path))

    for handle in _iter_parts(document):
        paragraphs = list(handle.root.iter(qn("w:p")))
        for index, paragraph in enumerate(paragraphs):
            unit_key = f"docx:{handle.part_id}:p:{index:05d}"
            target = translations.get(unit_key)
            if target is None:
                continue
            groups = _group_runs(paragraph)
            _rewrite_paragraph(paragraph, groups, target)
        handle.flush()

    apply_direction(document, rtl=target_rtl, lang_tag=lang_tag)
    for handle in _iter_parts(document):
        handle.flush()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


# ---------------------------------------------------------------------------
# ضبط اتجاه المستند حسب لغة الهدف
# ---------------------------------------------------------------------------
# ترتيب العناصر في OOXML مُلزَم بالمخطّط: إضافة عنصر في آخر الأب
# بتخالف المخطّط وWord ممكن يشتكي أو يتجاهله. لكل عنصر بنحدد الوسوم
# اللي المفروض ييجي قبلها.
_PPR_AFTER_BIDI = (
    "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap",
    "w:jc", "w:textDirection", "w:textAlignment", "w:outlineLvl",
    "w:rPr", "w:sectPr", "w:pPrChange",
)
_RPR_AFTER_RTL = ("w:cs", "w:em", "w:lang", "w:eastAsianLayout",
                  "w:specVanish", "w:oMath")
_TBLPR_AFTER_BIDIVISUAL = (
    "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
    "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd", "w:tblLayout",
    "w:tblCellMar", "w:tblLook",
)
_SECTPR_AFTER_BIDI = ("w:rtlGutter", "w:docGrid", "w:printerSettings",
                      "w:sectPrChange")


def _insert_ordered(parent, child, successors: tuple[str, ...]) -> None:
    """إدراج عنصر في موضعه الصحيح حسب مخطّط OOXML."""
    for tag in successors:
        found = parent.find(qn(tag))
        if found is not None:
            found.addprevious(child)
            return
    parent.append(child)


def _ensure(parent, tag: str, successors: tuple[str, ...]) -> None:
    if parent.find(qn(tag)) is None:
        _insert_ordered(parent, OxmlElement(tag), successors)


def _drop(parent, tag: str) -> None:
    element = parent.find(qn(tag))
    if element is not None:
        parent.remove(element)


def apply_direction(document, rtl: bool, lang_tag: str = "en-US") -> None:
    """ضبط اتجاه المستند كله حسب لغة الهدف.

    النقطة دي بتفشل فيها أغلب أدوات الترجمة: بتستبدل النص وبس، فيطلع
    مستند إنجليزي بتخطيط عربي (أو العكس) — المحاذاة في الناحية الغلط،
    الجداول معكوسة، والترقيم مقلوب.

    بيشتغل في الاتجاهين: عربي→إنجليزي بيشيل خصائص RTL،
    وإنجليزي→عربي بيضيفها.
    """
    for handle in _iter_parts(document):
        root = handle.root

        # 1) اتجاه الفقرات + المحاذاة
        for ppr in root.iter(qn("w:pPr")):
            if rtl:
                _ensure(ppr, "w:bidi", _PPR_AFTER_BIDI)
            else:
                _drop(ppr, "w:bidi")

            jc = ppr.find(qn("w:jc"))
            if jc is not None:
                value = jc.get(qn("w:val"))
                # "بداية السطر" بتنقلب مع الاتجاه: يمين في العربي،
                # يسار في الإنجليزي. المحاذاة للوسط والضبط مابيتغيروش.
                flip = {"right": "left", "left": "right",
                        "end": "start", "start": "end"}
                wanted_start = "right" if rtl else "left"
                if value in flip and value != wanted_start and value in ("left", "right"):
                    jc.set(qn("w:val"), wanted_start)

        # 2) اتجاه الحروف داخل الـ runs
        for rpr in root.iter(qn("w:rPr")):
            if rtl:
                _ensure(rpr, "w:rtl", _RPR_AFTER_RTL)
            else:
                _drop(rpr, "w:rtl")

            lang = rpr.find(qn("w:lang"))
            if lang is not None:
                lang.set(qn("w:val"), lang_tag)
                if rtl:
                    lang.set(qn("w:bidi"), lang_tag)
                else:
                    lang.attrib.pop(qn("w:bidi"), None)
                lang.attrib.pop(qn("w:eastAsia"), None)

        # 3) اتجاه الجداول (ترتيب الأعمدة بينقلب)
        for tblpr in root.iter(qn("w:tblPr")):
            if rtl:
                _ensure(tblpr, "w:bidiVisual", _TBLPR_AFTER_BIDIVISUAL)
            else:
                _drop(tblpr, "w:bidiVisual")

        # 4) اتجاه المقاطع (sections)
        for sectpr in root.iter(qn("w:sectPr")):
            if rtl:
                _ensure(sectpr, "w:bidi", _SECTPR_AFTER_BIDI)
            else:
                _drop(sectpr, "w:bidi")

        handle.flush()

    # 5) الاتجاه الافتراضي في أنماط المستند (styles.xml)
    _apply_styles_direction(document, rtl, lang_tag)


def _apply_styles_direction(document, rtl: bool, lang_tag: str) -> None:
    """ضبط الاتجاه في الأنماط وافتراضيات المستند.

    أغلب الفقرات في مستند Word مالهاش `w:pPr` أصلًا — بتورث الاتجاه
    من افتراضيات المستند. فلو عدّلنا الفقرات الصريحة بس، أغلب
    المستند هيفضل بالاتجاه القديم.
    """
    try:
        styles = document.styles.element
    except Exception:
        return

    # 1) افتراضيات المستند — دي اللي بتغطّي الفقرات اللي مالهاش إعداد صريح
    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None and rtl:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)

    if defaults is not None:
        for wrapper_tag, inner_tag, mark_tag, successors in (
            ("w:pPrDefault", "w:pPr", "w:bidi", _PPR_AFTER_BIDI),
            ("w:rPrDefault", "w:rPr", "w:rtl", _RPR_AFTER_RTL),
        ):
            wrapper = defaults.find(qn(wrapper_tag))
            if wrapper is None:
                if not rtl:
                    continue
                wrapper = OxmlElement(wrapper_tag)
                defaults.append(wrapper)

            inner = wrapper.find(qn(inner_tag))
            if inner is None:
                if not rtl:
                    continue
                inner = OxmlElement(inner_tag)
                wrapper.append(inner)

            if rtl:
                _ensure(inner, mark_tag, successors)
            else:
                _drop(inner, mark_tag)

    # 2) الأنماط المسمّاة
    for ppr in styles.iter(qn("w:pPr")):
        if rtl:
            _ensure(ppr, "w:bidi", _PPR_AFTER_BIDI)
        else:
            _drop(ppr, "w:bidi")

    for rpr in styles.iter(qn("w:rPr")):
        if rtl:
            _ensure(rpr, "w:rtl", _RPR_AFTER_RTL)
        else:
            _drop(rpr, "w:rtl")
        lang = rpr.find(qn("w:lang"))
        if lang is not None:
            lang.set(qn("w:val"), lang_tag)
