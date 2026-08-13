"""اختبار الاتجاه العكسي: إنجليزي → عربي.

الحاجات اللي بتتشقلب مع الاتجاه:
  1. اتجاه المستند لازم يبقى RTL بدل LTR.
  2. فحص «غير مترجم» لازم يقيس الكتابة الصح (لو فضل بيدوّر على عربي
     في الهدف، هيعلّم على كل مقطع مترجم صح).
  3. تعليمات النموذج لازم تتبدّل — مشاكل الاتجاهين مختلفة.
  4. تقدير التكلفة لازم يستخدم كثافة توكن كل لغة.
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from sqlalchemy import select  # noqa: E402

from app.core.db import SessionLocal, init_db  # noqa: E402
from app.models import Project, Segment  # noqa: E402
from app.tools.translator import pipeline  # noqa: E402
from app.tools.translator.costing import estimate_project  # noqa: E402
from app.tools.translator.engine import (  # noqa: E402
    BatchResult,
    SegmentInput,
    Usage,
    validate_translation,
)
from app.tools.translator.prompts import build_system_prompt  # noqa: E402

SAMPLE = Path("storage/samples/contract_en.source.docx")

# ترجمة عربية وهمية لكل جملة إنجليزية في العيّنة
FAKE_AR = {
    "Consultancy Services Agreement": "اتفاقية تقديم خدمات استشارية",
    "This Agreement is made on the fifteenth day of March.": "أُبرمت هذه الاتفاقية في اليوم الخامس عشر من مارس.",
    "Article 1. The First Party shall provide the agreed services.": "المادة 1. يلتزم الطرف الأول بتقديم الخدمات المتفق عليها.",
    "Article 2. The contract value is 150,000 USD.": "المادة 2. تبلغ قيمة العقد 150,000 دولار أمريكي.",
    "Item": "البند",
    "Description": "الوصف",
    "Value": "القيمة",
    "Fees": "الأتعاب",
    "Monthly consultancy fees": "أتعاب الاستشارة الشهرية",
    "12,500": "12,500",
}


class ArabicEchoEngine:
    """محرّك وهمي بيرجّع عربي — عشان نختبر الاتجاه العكسي بدون API."""

    name = "ar-echo"
    model = "ar-echo"
    production = False

    def translate(self, segments, context_before="", context_after=""):
        result = BatchResult(usage=Usage(calls=1))
        for segment in segments:
            plain = segment.source.strip()
            result.translations[segment.id] = FAKE_AR.get(
                plain, f"نص مترجم للعربية: {plain}"
            )
        return result


def build_english_docx(path: Path) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("Consultancy Services Agreement", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("This Agreement is made on the fifteenth day of March.")
    doc.add_paragraph("Article 1. The First Party shall provide the agreed services.")
    doc.add_paragraph("Article 2. The contract value is 150,000 USD.")

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for col, value in enumerate(["Item", "Description", "Value"]):
        table.cell(0, col).text = value
    for col, value in enumerate(["Fees", "Monthly consultancy fees", "12,500"]):
        table.cell(1, col).text = value

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> int:
    failures: list[str] = []
    init_db()
    build_english_docx(SAMPLE)

    # ---------- 1) التعليمات بتتبدّل مع الاتجاه ----------
    print("=== 1) تعليمات النموذج ===")
    ar_en = build_system_prompt(source_lang="ar", target_lang="en", domain="legal")
    en_ar = build_system_prompt(source_lang="en", target_lang="ar", domain="legal")
    print(f"  عربي→إنجليزي: {len(ar_en):,} حرف")
    print(f"  إنجليزي→عربي: {len(en_ar):,} حرف")

    if "Arabic-to-English specifics" not in ar_en:
        failures.append("تعليمات عربي→إنجليزي مش موجودة")
    if "English-to-Arabic specifics" not in en_ar:
        failures.append("تعليمات إنجليزي→عربي مش موجودة")
    if "Grammatical agreement" not in en_ar:
        failures.append("قواعد المطابقة النحوية العربية ناقصة")
    if "Grammatical agreement" in ar_en:
        failures.append("قواعد الاتجاه العكسي ظهرت في الاتجاه الغلط")
    print("  ✓ التعليمات بتتبدّل حسب الاتجاه")

    # ---------- 2) فحص الجودة محايد للاتجاه ----------
    print("\n=== 2) فحوصات الجودة ===")
    cases = [
        ("ar", "en", "المادة 1. يلتزم الطرف الأول.", "Article 1. The First Party undertakes.", False),
        ("ar", "en", "المادة 1. يلتزم الطرف الأول.", "المادة 1. يلتزم الطرف الأول.", True),
        ("en", "ar", "Article 1. The First Party undertakes.", "المادة 1. يلتزم الطرف الأول.", False),
        ("en", "ar", "Article 1. The First Party undertakes.", "Article 1. The First Party undertakes.", True),
    ]
    for source_lang, target_lang, source, target, expect_flag in cases:
        problems = validate_translation(source, target, source_lang, target_lang)
        flagged = "untranslated" in problems
        status = "✓" if flagged == expect_flag else "✗"
        print(f"  {status} {source_lang}→{target_lang} "
              f"{'(غير مترجم)' if expect_flag else '(مترجم)'}: {problems or 'سليم'}")
        if flagged != expect_flag:
            failures.append(
                f"فحص {source_lang}→{target_lang} غلط: متوقع flag={expect_flag}"
            )

    # ---------- 3) الخط الكامل إنجليزي→عربي ----------
    print("\n=== 3) الخط الكامل (إنجليزي → عربي) ===")
    db = SessionLocal()
    project = None
    try:
        project = Project(
            name="اختبار عكسي — إنجليزي إلى عربي",
            source_lang="en",
            target_lang="ar",
            domain="legal",
        )
        db.add(project)
        db.flush()

        file = pipeline.ingest(db, project, SAMPLE, SAMPLE.name)
        db.commit()

        stats = pipeline.extract_and_segment(db, file)
        db.commit()
        print(f"  استخراج: {stats.units} وحدة · {stats.segments} مقطع")

        translated = pipeline.translate_file(
            db, file, ArabicEchoEngine(), use_memory=False
        )
        db.commit()
        print(f"  ترجمة: {translated.translated} مقطع · "
              f"أعلام={translated.flagged}")

        if translated.flagged > 2:
            flagged_segments = db.execute(
                select(Segment).where(
                    Segment.file_id == file.id, Segment.qa_flags != "[]"
                )
            ).scalars().all()
            for segment in flagged_segments[:4]:
                print(f"    ⚠ {segment.location}: {segment.qa_flags}")
            failures.append(
                f"{translated.flagged} مقطع عليه أعلام رغم إن الترجمة سليمة — "
                "الفحص لسه مربوط باتجاه واحد"
            )

        # ---------- 4) التصدير باتجاه عربي ----------
        output = pipeline.export_file(db, file)
        db.commit()
        print(f"\n=== 4) التصدير: {output.name} ===")

        if not output.name.endswith(".ar.docx"):
            failures.append(f"امتداد المخرَج مش متوقع: {output.name}")

        doc = Document(str(output))
        text = "\n".join(p.text for p in doc.paragraphs)
        text += "\n" + "\n".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells
        )

        arabic = sum(1 for ch in text if "؀" <= ch <= "ۿ")
        latin = sum(1 for ch in text if "a" <= ch.lower() <= "z")
        print(f"  حروف عربية={arabic} · حروف لاتينية={latin}")
        if arabic < 100:
            failures.append(f"النص العربي قليل في المخرَج: {arabic} حرف")

        # الاتجاه لازم يكون اتقلب لـ RTL
        bidi = len(list(doc.element.iter(qn("w:bidi"))))
        rtl = len(list(doc.element.iter(qn("w:rtl"))))
        print(f"  عناصر اتجاه عربي في المتن: bidi={bidi} rtl={rtl}")
        if bidi == 0:
            failures.append("الاتجاه مااتقلبش لـ RTL — المستند هيطلع بتخطيط إنجليزي")

        # الأهم: افتراضيات المستند — دي اللي بتغطّي الفقرات اللي
        # مالهاش إعداد صريح، وهي أغلب المستند عادةً
        styles = doc.styles.element
        defaults = styles.find(qn("w:docDefaults"))
        default_bidi = default_rtl = 0
        if defaults is not None:
            default_bidi = len(list(defaults.iter(qn("w:bidi"))))
            default_rtl = len(list(defaults.iter(qn("w:rtl"))))
        print(f"  افتراضيات المستند: bidi={default_bidi} rtl={default_rtl}")
        if default_bidi == 0:
            failures.append(
                "افتراضيات المستند مش RTL — الفقرات اللي مالهاش إعداد صريح "
                "هتفضل بتخطيط إنجليزي"
            )
        if default_rtl == 0:
            failures.append("افتراضيات الحروف مش RTL")

        # اتجاه الجداول
        tbl_bidi = len(list(doc.element.iter(qn("w:bidiVisual"))))
        print(f"  جداول باتجاه عربي: {tbl_bidi}")
        if doc.tables and tbl_bidi == 0:
            failures.append("ترتيب أعمدة الجدول مااتقلبش")

        # الملف لازم يفتح في Word من غير كسر — نتأكد إنه يتقرا تاني
        from app.tools.translator.formats import docx_fmt

        reread = docx_fmt.extract(output)
        print(f"  الملف يتقرا تاني: {len(reread.units)} وحدة")
        if len(reread.units) < stats.units - 1:
            failures.append("الملف المصدَّر اتكسر — الاستخراج منه ناقص")

    finally:
        if project is not None:
            db.query(Project).filter(Project.id == project.id).delete()
            db.commit()
        db.close()

    # ---------- 5) التقدير حسب كثافة اللغة ----------
    print("\n=== 5) تقدير التكلفة حسب اتجاه الترجمة ===")
    ar_est = estimate_project(1000, 6000, 4, 60, source_lang="ar", target_lang="en")
    en_est = estimate_project(1000, 6000, 4, 60, source_lang="en", target_lang="ar")
    ar_cost = ar_est.options[0]["cost_usd"]
    en_cost = en_est.options[0]["cost_usd"]
    print(f"  نفس عدد الحروف (6000):")
    print(f"    عربي→إنجليزي: ${ar_cost:.4f}  ({ar_est.input_tokens:,} توكن دخل)")
    print(f"    إنجليزي→عربي: ${en_cost:.4f}  ({en_est.input_tokens:,} توكن دخل)")

    # العربية أكثف بالتوكن، فنفس عدد الحروف بيكلف أكتر لما تكون هي المصدر
    if ar_est.input_tokens <= en_est.input_tokens:
        failures.append(
            "التقدير مش بياخد كثافة اللغة في الحسبان — "
            "العربية المفروض توكناتها أكتر لنفس عدد الحروف"
        )
    else:
        print("  ✓ التقدير بيفرّق بين كثافة اللغتين")

    print("\n" + "=" * 62)
    if failures:
        print(f"فشل: {len(failures)}")
        for item in failures:
            print(f"  ✗ {item}")
        return 1
    print("نجح: الاتجاه العكسي (إنجليزي → عربي) سليم ✓")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
