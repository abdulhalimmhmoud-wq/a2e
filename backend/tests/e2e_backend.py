"""اختبار شامل لخط الأنابيب الكامل بدون أي نداء API.

بيغطّي: إنشاء مشروع → رفع ملف → استخراج → تقسيم → ترجمة →
ذاكرة الترجمة → تعديل يدوي → انتشار → تصدير → تقرير التكلفة.
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from sqlalchemy import select  # noqa: E402

from app.core.db import SessionLocal, init_db  # noqa: E402
from app.models import Project, Segment, SourceFile  # noqa: E402
from app.tools.translator import pipeline, tm  # noqa: E402
from app.tools.translator.engine import EchoEngine  # noqa: E402
from app.tools.translator.formats.base import strip_tags  # noqa: E402

SAMPLE = Path("storage/samples/contract_ar.docx")


def main() -> int:
    failures: list[str] = []
    init_db()
    db = SessionLocal()
    project = None

    try:
        # ---------- 1) مشروع + رفع ----------
        project = Project(
            name="اختبار شامل — عقد استشارات",
            domain="legal",
            model="claude-sonnet-5",
        )
        db.add(project)
        db.flush()

        if not SAMPLE.exists():
            print("!! شغّل make_sample_docx.py الأول")
            return 1

        file = pipeline.ingest(db, project, SAMPLE, SAMPLE.name)
        db.commit()
        print(f"=== 1) الرفع: {file.original_filename} ({file.fmt}) ===")

        # ---------- 2) الاستخراج والتقسيم ----------
        stats = pipeline.extract_and_segment(db, file)
        db.commit()
        print(f"\n=== 2) الاستخراج ===")
        print(f"  وحدات={stats.units} مقاطع={stats.segments} "
              f"كلمات={stats.words} حروف={stats.chars} صفحات={stats.pages}")

        if stats.segments < 15:
            failures.append(f"عدد المقاطع قليل: {stats.segments}")
        if stats.words < 40:
            failures.append(f"عدد الكلمات قليل: {stats.words}")

        translatable = db.execute(
            select(Segment).where(
                Segment.file_id == file.id, Segment.is_translatable.is_(True)
            )
        ).scalars().all()
        print(f"  قابل للترجمة: {len(translatable)}/{stats.segments}")

        # ---------- 3) الترجمة ----------
        result = pipeline.translate_file(db, file, EchoEngine(), use_memory=True)
        db.commit()
        print(f"\n=== 3) الترجمة (محرّك تجريبي) ===")
        print(f"  مترجَم={result.translated} من_الذاكرة={result.from_memory} "
              f"فشل={result.failed} عليه_أعلام={result.flagged}")

        # المقاطع بتتغطّى من مصدرين: المحرّك + ذاكرة الترجمة.
        # لو الذاكرة فيها نتائج من تشغيلة سابقة، عدّاد المحرّك بيقل — وده صح.
        covered = result.translated + result.from_memory
        if covered < len(translatable):
            failures.append(
                f"مقاطع ماتترجمتش: {len(translatable) - covered} "
                f"(محرّك={result.translated} ذاكرة={result.from_memory})"
            )
        if result.failed:
            failures.append(f"مقاطع فشلت: {result.failed}")

        progress = pipeline.file_progress(db, file.id)
        print(f"  الحالة: {progress}")

        # ---------- 4) حاجز الذاكرة + الاعتماد ----------
        print(f"\n=== 4) ذاكرة الترجمة ===")

        # أ) مخرجات المحرّك التجريبي ممنوعة من الذاكرة.
        #    لو دخلت، هترجع في كل مشروع حقيقي بعد كده وتفسد الترجمة.
        for segment in translatable[:5]:
            db.refresh(segment)
            pipeline.approve_segment(db, segment)
        db.commit()

        echo_segment = translatable[0]
        db.refresh(echo_segment)
        leaked = tm.lookup_exact(db, echo_segment.source_text, "ar", "en", "legal")
        print(f"  مخرجات تجريبية دخلت الذاكرة: {'نعم' if leaked else 'لا'}")
        if leaked is not None:
            failures.append(
                "مخرجات المحرّك التجريبي دخلت ذاكرة الترجمة — هتلوّث المشاريع الحقيقية"
            )

        # ب) التعديل البشري لازم يتخزّن في الذاكرة
        human_segment = translatable[6]
        db.refresh(human_segment)
        human_segment.target_text = "A binding consultancy agreement."
        human_segment.origin = "human"
        human_segment.engine_model = ""
        human_segment.edited_by_human = True
        db.flush()
        pipeline.approve_segment(db, human_segment)
        db.commit()

        stored = tm.lookup_exact(db, human_segment.source_text, "ar", "en", "legal")
        print(f"  تعديل بشري اتخزّن في الذاكرة: {'نعم' if stored else 'لا'}")
        if stored is None:
            failures.append("التعديل البشري ماتخزّنش في ذاكرة الترجمة")

        # ---------- 5) التعديل والانتشار ----------
        print(f"\n=== 5) التعديل والانتشار ===")

        # نجهّز حالة تكرار: مقطعين بنفس النص المصدر
        repeated = db.execute(
            select(Segment).where(Segment.file_id == file.id).limit(2)
        ).scalars().all()
        repeated[1].source_text = repeated[0].source_text
        repeated[1].source_hash = repeated[0].source_hash
        repeated[1].target_text = repeated[0].target_text
        db.flush()

        new_text = repeated[0].target_text + " (REVISED)"
        plan = tm.plan_propagation(db, repeated[0], new_text)
        applied = tm.apply_propagation(db, plan.auto)
        repeated[0].target_text = new_text
        repeated[0].edited_by_human = True
        db.commit()

        print(f"  مطابقة تامة: تلقائي={len(plan.auto)} محتاج_موافقة={len(plan.needs_review)}")
        print(f"  اتطبّق تلقائيًا: {applied}")
        if applied < 1:
            failures.append("الانتشار على المطابقة التامة مااشتغلش")

        db.refresh(repeated[1])
        if "REVISED" not in repeated[1].target_text:
            failures.append("المقطع المكرر مااتحدّثش")

        # انتشار على مستوى المصطلح
        term_segment = translatable[3]
        db.refresh(term_segment)
        before = term_segment.target_text
        changed = before.replace("EN", "ENG", 1) if "EN" in before else before + " X"
        term_plan = tm.plan_propagation(db, term_segment, changed)
        print(f"  مستوى المصطلح: مرشّحين={len(term_plan.needs_review)}")

        # ---------- 6) التصدير ----------
        output = pipeline.export_file(db, file)
        db.commit()
        print(f"\n=== 6) التصدير ===")
        print(f"  {output.name} ({output.stat().st_size:,} بايت)")

        if not output.exists() or output.stat().st_size < 5000:
            failures.append("ملف التصدير مفقود أو صغير بشكل مريب")

        # التحقق: النص المترجم فعلًا جوّه الملف
        from docx import Document

        doc = Document(str(output))
        body = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        combined = body + "\n" + table_text

        if "[EN]" not in combined:
            failures.append("الترجمة مش موجودة في الملف المصدَّر")
        if "REVISED" not in combined:
            failures.append("التعديل اليدوي مش موجود في الملف المصدَّر")

        arabic_left = sum(1 for c in combined if "؀" <= c <= "ۿ")
        print(f"  حروف عربية متبقية في المخرجات: {arabic_left}")

        # ---------- 7) التكلفة ----------
        from app.tools.translator.costing import estimate_project

        estimate = estimate_project(
            words=file.word_count,
            chars=file.char_count,
            pages=file.page_count,
            segments=file.segment_count,
        )
        print(f"\n=== 7) تقدير التكلفة ===")
        print(f"  {estimate.words} كلمة · {estimate.pages} صفحة · "
              f"~{estimate.input_tokens:,} توكن إدخال")
        for option in estimate.options:
            print(f"    {option['label']:20} ${option['cost_usd']:.4f}  "
                  f"(دفعات: ${option['cost_usd_batch']:.4f})  "
                  f"${option['cost_per_page']:.4f}/صفحة")

        if not estimate.options:
            failures.append("التقدير مارجّعش أي خيار")

    finally:
        if project is not None:
            db.query(Project).filter(Project.id == project.id).delete()
            db.commit()
        db.close()

    print("\n" + "=" * 62)
    if failures:
        print(f"فشل: {len(failures)} مشكلة")
        for item in failures:
            print(f"  ✗ {item}")
        return 1
    print("نجح: خط الأنابيب الكامل سليم ✓")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
