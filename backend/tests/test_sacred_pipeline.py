"""اختبار مسار النص المقدّس كامل: استخراج ← قفل ← ترجمة ← تصدير.

المحرّك هنا وهمي وبيسجّل كل مقطع اتبعتله. الشرط الأساسي إن الآيات
والأحاديث **مايوصلوش** للمحرّك أصلًا، مش إن ترجمتهم تتصلّح بعدين.
مفيش أي نداء API.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from sqlalchemy import select  # noqa: E402

from app.core.db import SessionLocal, init_db  # noqa: E402
from app.models import Project, Segment  # noqa: E402
from app.tools.translator import pipeline  # noqa: E402
from app.tools.translator.engine import BatchResult, Usage  # noqa: E402

WORK = Path("storage/samples/sacred")

PARAGRAPHS = [
    "عقد مرابحة إسلامية",
    "المادة 1. يلتزم البائع بتسليم البضاعة في الموعد المتفق عليه.",
    "قال تعالى: ﴿وَأَحَلَّ اللَّهُ الْبَيْعَ وَحَرَّمَ الرِّبَا﴾ [البقرة: 275]",
    "المادة 2. لا يجوز اشتراط أي زيادة مقابل التأجيل.",
    "قال رسول الله صلى الله عليه وسلم: «المسلمون على شروطهم» رواه البخاري.",
    "المادة 3. تخضع هذه الاتفاقية لأحكام هيئة الرقابة الشرعية.",
    "وهو ما يفهم من الحديث النبوي الوارد في هذا الباب.",
]


class RecordingEngine:
    """محرّك وهمي بيسجّل اللي اتبعتله."""

    model = "fake-model"
    production = False

    def __init__(self) -> None:
        self.seen: list[str] = []

    def translate(self, segments, context_before="", context_after=""):
        self.seen.extend(s.source for s in segments)
        return BatchResult(
            translations={s.id: f"[EN] {s.source}" for s in segments},
            usage=Usage(),
        )


def build_docx(path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    doc = Document()
    for line in PARAGRAPHS:
        paragraph = doc.add_paragraph(line)
        paragraph._p.get_or_add_pPr().append(OxmlElement("w:bidi"))  # noqa: SLF001
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> int:  # noqa: C901
    failures: list[str] = []
    init_db()

    source = WORK / "murabaha_ar.docx"
    build_docx(source)

    db = SessionLocal()
    project = None
    try:
        project = Project(
            name="__اختبار النص المقدّس__",
            source_lang="ar",
            target_lang="en",
            domain="legal",
        )
        db.add(project)
        db.commit()

        # ---------- 1) الاستخراج والكشف ----------
        file = pipeline.ingest(db, project, source, source.name)
        db.commit()
        stats = pipeline.extract_and_segment(db, file)
        db.commit()

        print("=== 1) الاستخراج ===")
        print(f"  مقاطع: {stats.segments} · مقفول: {stats.sacred_locked} · "
              f"معلَّم: {stats.sacred_flagged}")

        segments = db.execute(
            select(Segment)
            .where(Segment.file_id == file.id)
            .order_by(Segment.order_index)
        ).scalars().all()

        print("\n=== 2) حالة كل مقطع ===")
        for segment in segments:
            mark = "🔒" if segment.is_locked else "  "
            print(f"  {mark} {segment.source_text[:52]}")
            if segment.notes:
                print(f"       {segment.notes[:66]}")

        locked = [s for s in segments if s.is_locked]
        if stats.sacred_locked < 2:
            failures.append(
                f"المفروض يتقفل على الأقل مقطعين (آية + حديث)، "
                f"اتقفل {stats.sacred_locked}"
            )

        # البنود القانونية العادية ماينفعش تتقفل
        for segment in segments:
            if segment.is_locked and segment.source_text.startswith("المادة"):
                failures.append(
                    f"بند قانوني اتقفل بالغلط: {segment.source_text[:44]}"
                )

        # ---------- 3) الترجمة ----------
        engine = RecordingEngine()
        pipeline.translate_file(db, file, engine, use_memory=False)
        db.commit()

        print("\n=== 3) اللي وصل للمحرّك ===")
        for text in engine.seen:
            print(f"    {text[:56]}")

        sent = " ".join(engine.seen)
        if "﴿" in sent or "وَأَحَلَّ" in sent:
            failures.append("آية قرآنية اتبعتت للمحرّك — ده بالظبط اللي المفروض يتمنع")
        if "المسلمون على شروطهم" in sent:
            failures.append("حديث اتبعت للمحرّك")
        if not any("المادة 1" in text for text in engine.seen):
            failures.append("بند قانوني عادي مااتبعتش — القفل زايد عن اللزوم")

        # ---------- 4) الأعلام بعد الترجمة ----------
        db.refresh(file)
        segments = db.execute(
            select(Segment)
            .where(Segment.file_id == file.id)
            .order_by(Segment.order_index)
        ).scalars().all()

        print("\n=== 4) الأعلام بعد الترجمة ===")
        sacred_flags = 0
        for segment in segments:
            flags = segment.qa_flags
            if "quran" in flags or "hadith" in flags:
                sacred_flags += 1
                state = "مقفول" if segment.is_locked else "مترجَم ومعلَّم"
                print(f"  {state}: {segment.source_text[:42]}")
                print(f"      {flags}")

        if sacred_flags < 3:
            failures.append(
                f"أعلام النص المقدّس ضاعت بعد الترجمة: فاضل {sacred_flags}"
            )

        # المقطع المرجَّح (مش مقفول) لازم يكون اتترجم ولسه معلَّم
        likely = [
            s for s in segments
            if not s.is_locked and ("quran_likely" in s.qa_flags
                                    or "hadith_likely" in s.qa_flags)
        ]
        if not likely:
            failures.append(
                "المقطع المرجَّح فقد علامته بعد الترجمة — فحوصات ما بعد "
                "الترجمة داست عليها"
            )
        elif not likely[0].target_text:
            failures.append("المقطع المرجَّح مااتترجمش رغم إنه مش مقفول")

        # المقفول لازم يفضل فاضي
        for segment in locked:
            if segment.target_text:
                failures.append(
                    f"مقطع مقفول اتكتب فيه ترجمة: {segment.source_text[:40]}"
                )

        # ---------- 5) إعادة التشغيل ماتدوسش على المقفول ----------
        engine2 = RecordingEngine()
        pipeline.translate_file(db, file, engine2, use_memory=False)
        db.commit()
        print(f"\n=== 5) إعادة التشغيل: اتبعت {len(engine2.seen)} مقطع ===")
        if any("﴿" in text for text in engine2.seen):
            failures.append("إعادة التشغيل بعتت الآية للمحرّك")

        # ---------- 6) التصدير ----------
        exported = pipeline.export_file(db, file)
        db.commit()
        print(f"\n=== 6) التصدير: {Path(exported).name} ===")
        if not Path(exported).exists():
            failures.append("التصدير فشل مع وجود مقاطع مقفولة")
        else:
            from docx import Document

            text = "\n".join(p.text for p in Document(exported).paragraphs)
            # الآية لازم تفضل بنصها العربي في المخرجات
            if "وَأَحَلَّ" not in text and "أحل" not in text:
                failures.append("الآية اختفت من الملف المصدَّر")
            else:
                print("  ✓ الآية اتحافظ عليها بنصها العربي")

    finally:
        if project is not None:
            import shutil

            folder = pipeline.project_dir(project.id)
            db.query(Project).filter(Project.id == project.id).delete()
            db.commit()
            shutil.rmtree(folder, ignore_errors=True)
        db.close()

    print("\n" + "=" * 62)
    if failures:
        print(f"فشل: {len(failures)}")
        for item in failures:
            print(f"  ✗ {item}")
        return 1
    print("نجح: الآيات والأحاديث مابتوصلش للمحرّك ✓")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
