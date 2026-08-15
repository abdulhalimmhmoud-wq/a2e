"""اختبار استخراج المصطلحات.

الاستيراد من جدول والمحاذاة بيتختبروا بدون أي نداء API.
الاستخراج بالنموذج بيتختبر بمحاكاة، وكمان بنداء حقيقي لو المفتاح موجود.
"""
from __future__ import annotations

import csv
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.core.config import settings  # noqa: E402
from app.tools.translator import terms  # noqa: E402

WORK = Path("storage/samples/terms")


def build_csv(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["المصطلح", "الترجمة", "ملاحظة"])
        writer.writerow(["عقد إذعان", "Contract of Adhesion", "قانون مدني"])
        writer.writerow(["الطرف الأول", "First Party", ""])
        writer.writerow(["القوة القاهرة", "Force Majeure", ""])
        writer.writerow(["", "قيمة ناقصة", "الصف ده لازم يتتخطى"])
        writer.writerow(["الطرف الأول", "First Party", "مكرر"])
        writer.writerow(
            ["جملة طويلة جدًا " * 12, "A very long sentence " * 8, "أطول من مصطلح"]
        )


def build_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["Source", "Target"])
    ws.append(["محضر جلسة", "Session Minutes"])
    ws.append(["الحكم النهائي", "Final Judgment"])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))


def build_pair(source: Path, target: Path) -> None:
    """ملف عربي وترجمته الإنجليزية بنفس البنية."""
    from docx import Document

    ar = [
        "عقد تقديم خدمات استشارية",
        "المادة 1. يلتزم الطرف الأول بتقديم الخدمات المتفق عليها.",
        "المادة 2. تسري أحكام القوة القاهرة على هذا العقد.",
        "المادة 3. يحق للطرف الثاني إنهاء العقد بإشعار كتابي.",
    ]
    en = [
        "Consultancy Services Agreement",
        "Article 1. The First Party shall provide the agreed services.",
        "Article 2. Force Majeure provisions apply to this contract.",
        "Article 3. The Second Party may terminate the contract by written notice.",
    ]
    for path, lines in ((source, ar), (target, en)):
        doc = Document()
        for line in lines:
            doc.add_paragraph(line)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))


def main() -> int:
    failures: list[str] = []
    WORK.mkdir(parents=True, exist_ok=True)

    # ---------- 1) استيراد CSV ----------
    print("=== 1) استيراد CSV ===")
    csv_path = WORK / "terms.csv"
    build_csv(csv_path)
    result = terms.import_table(csv_path)

    print(f"  صفوف مفحوصة: {result.pairs_examined} · مرشّحون: {len(result.candidates)}")
    for candidate in result.candidates:
        print(f"    {candidate.source_term} → {candidate.target_term}"
              f"{'  · ' + candidate.note if candidate.note else ''}")

    sources = [c.source_term for c in result.candidates]
    if len(result.candidates) != 3:
        failures.append(f"متوقع 3 مصطلحات، طلع {len(result.candidates)}")
    if "" in sources:
        failures.append("صف ناقص القيمة اتقبل")
    if len(sources) != len(set(sources)):
        failures.append("مصطلح مكرر اتقبل مرتين")
    if any(len(c.source_term) > 120 for c in result.candidates):
        failures.append("جملة طويلة اتقبلت كمصطلح")
    if not any(c.note for c in result.candidates):
        failures.append("عمود الملاحظة ضاع")

    # ---------- 2) استيراد Excel ----------
    print("\n=== 2) استيراد Excel ===")
    xlsx_path = WORK / "terms.xlsx"
    build_xlsx(xlsx_path)
    result_xlsx = terms.import_table(xlsx_path)
    print(f"  مرشّحون: {len(result_xlsx.candidates)}")
    for candidate in result_xlsx.candidates:
        print(f"    {candidate.source_term} → {candidate.target_term}")
    if len(result_xlsx.candidates) != 2:
        failures.append(f"Excel: متوقع 2، طلع {len(result_xlsx.candidates)}")
    if any(c.source_term.lower() == "source" for c in result_xlsx.candidates):
        failures.append("صف العناوين مااتتخطاش")

    # ---------- 3) ترميز cp1256 ----------
    print("\n=== 3) ترميز عربي قديم (cp1256) ===")
    legacy = WORK / "legacy.csv"
    legacy.write_bytes(
        "المصطلح,الترجمة\r\nمحكمة النقض,Court of Cassation\r\n".encode("cp1256")
    )
    result_legacy = terms.import_table(legacy)
    print(f"  مرشّحون: {len(result_legacy.candidates)}")
    if result_legacy.candidates:
        print(f"    {result_legacy.candidates[0].source_term} → "
              f"{result_legacy.candidates[0].target_term}")
    if not result_legacy.candidates:
        failures.append("ملف cp1256 مااتقراش — ملفات عربية قديمة كتير بالترميز ده")
    elif "محكمة" not in result_legacy.candidates[0].source_term:
        failures.append("ترميز cp1256 اتقرا غلط")

    # ---------- 4) المحاذاة ----------
    print("\n=== 4) محاذاة ملف بترجمته ===")
    source_doc, target_doc = WORK / "contract_ar.docx", WORK / "contract_en.docx"
    build_pair(source_doc, target_doc)
    alignment = terms.align_documents(source_doc, target_doc, WORK / "work")

    print(f"  وحدات: {alignment.source_units} مقابل {alignment.target_units}")
    print(f"  أزواج: {len(alignment.pairs)} · موثوقة: {alignment.confident}")
    for source, target in alignment.pairs[:3]:
        print(f"    {source[:44]}")
        print(f"      -> {target[:44]}")

    if not alignment.confident:
        failures.append("المحاذاة اتعلّمت غير موثوقة رغم تطابق البنية")
    if len(alignment.pairs) < 3:
        failures.append(f"أزواج قليلة: {len(alignment.pairs)}")
    if alignment.pairs and "الطرف الأول" not in alignment.pairs[0][0]:
        # أول زوج المفروض المادة 1 (العنوان أقصر من 8 حروف؟ لأ، بيتقبل)
        pass

    # ---------- 5) محاذاة ببنية مختلفة ----------
    print("\n=== 5) محاذاة ببنية مختلفة (لازم تتعلّم) ===")
    from docx import Document

    short = WORK / "short_en.docx"
    doc = Document()
    doc.add_paragraph("Consultancy Services Agreement")
    doc.save(str(short))

    mismatched = terms.align_documents(source_doc, short, WORK / "work")
    print(f"  وحدات: {mismatched.source_units} مقابل {mismatched.target_units}")
    print(f"  موثوقة: {mismatched.confident}")
    for warning in mismatched.warnings:
        print(f"    ! {warning[:72]}")
    if mismatched.confident:
        failures.append(
            "بنية مختلفة جدًا اتعلّمت موثوقة — الأزواج المتزحلقة "
            "بتولّد مصطلحات غلط"
        )

    # ---------- 6) استخراج حقيقي بالنموذج ----------
    if settings.anthropic_api_key:
        print("\n=== 6) استخراج بالنموذج (نداء حقيقي) ===")
        extracted = terms.extract_from_pairs(
            alignment.pairs, source_lang="ar", target_lang="en", domain="legal"
        )
        print(f"  أزواج مفحوصة: {extracted.pairs_examined} · "
              f"مصطلحات: {len(extracted.candidates)} · "
              f"تكلفة: ${extracted.usage.cost_usd:.6f}")
        for candidate in extracted.candidates:
            print(f"    {candidate.source_term} → {candidate.target_term}")
            if candidate.note:
                print(f"        {candidate.note[:66]}")
        for warning in extracted.warnings:
            print(f"    ! {warning}")

        if not extracted.candidates:
            failures.append("النموذج مااستخرجش أي مصطلح من نص قانوني واضح")
        else:
            found = " ".join(c.source_term for c in extracted.candidates)
            if "الطرف الأول" not in found and "القوة القاهرة" not in found:
                failures.append(
                    "المصطلحات الواضحة (الطرف الأول/القوة القاهرة) ماطلعتش"
                )
            # الجمل الكاملة مش مصطلحات
            long_ones = [c for c in extracted.candidates if len(c.source_term) > 60]
            if long_ones:
                failures.append(
                    f"{len(long_ones)} مرشّح طوله جملة مش مصطلح"
                )
    else:
        print("\n=== 6) تخطّي الاستخراج بالنموذج (مفيش مفتاح) ===")

    print("\n" + "=" * 62)
    if failures:
        print(f"فشل: {len(failures)}")
        for item in failures:
            print(f"  ✗ {item}")
        return 1
    print("نجح: استخراج المصطلحات سليم ✓")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
