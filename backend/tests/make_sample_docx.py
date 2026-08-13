"""توليد ملف Word عربي واقعي للاختبار.

بيحتوي على الحالات الصعبة اللي بتكسر أدوات الترجمة عادةً:
فقرات بتنسيق مختلط، جدول، ترويسة، رابط، قائمة مرقّمة، ونص RTL.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_rtl(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    ppr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)

    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def build(output: Path) -> None:
    doc = Document()

    # ترويسة
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "شركة المستقبل للاستشارات القانونية"
    _set_rtl(header_para)

    # عنوان
    title = doc.add_heading("عقد تقديم خدمات استشارية", level=1)
    _set_rtl(title)

    # فقرة بتنسيق مختلط — دي الحالة الصعبة
    p = doc.add_paragraph()
    _set_rtl(p)
    p.add_run("أُبرم هذا العقد في ")
    bold = p.add_run("اليوم الخامس عشر من شهر مارس")
    bold.bold = True
    p.add_run(" بين ")
    italic = p.add_run("الطرف الأول")
    italic.italic = True
    p.add_run(" والطرف الثاني، وفقًا لأحكام القانون المدني.")

    # فقرة فيها رابط
    p2 = doc.add_paragraph()
    _set_rtl(p2)
    p2.add_run("للاطلاع على اللائحة التنفيذية يُرجى زيارة ")
    _add_hyperlink(p2, "https://example.gov.sa/regulations", "الموقع الرسمي للوزارة")
    p2.add_run(" قبل التوقيع.")

    # قائمة مرقّمة فيها ترقيم قانوني (فخ للمقسّم)
    for text in [
        "المادة 1. يلتزم الطرف الأول بتقديم الخدمات المتفق عليها.",
        "المادة 2. تبلغ قيمة العقد 150.000 ريال سعودي.",
        "المادة 3. مدة العقد اثنا عشر شهرًا تبدأ من تاريخ التوقيع.",
    ]:
        item = doc.add_paragraph(text, style="List Number")
        _set_rtl(item)

    # فقرة متعددة الجُمل — لاختبار التقسيم
    p3 = doc.add_paragraph()
    _set_rtl(p3)
    p3.add_run(
        "يحق للطرف الثاني إنهاء العقد بإشعار كتابي. "
        "ويجب أن يُقدَّم الإشعار قبل ثلاثين يومًا على الأقل. "
        "هل يسري هذا الشرط على حالات القوة القاهرة؟ نعم، يسري."
    )

    # جدول
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    data = [
        ["البند", "الوصف", "القيمة"],
        ["الأتعاب", "أتعاب الاستشارة الشهرية", "12,500"],
        ["الضريبة", "ضريبة القيمة المضافة", "1,875"],
    ]
    for row_index, row in enumerate(data):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            for para in cell.paragraphs:
                _set_rtl(para)

    # تذييل
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.text = "صفحة 1 من 1 — وثيقة سرية"
    _set_rtl(footer_para)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"created: {output}")


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("storage/samples/contract_ar.docx")
    build(target)
