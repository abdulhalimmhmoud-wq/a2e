"""ترجمة حقيقية كاملة عبر الـ API — من الرفع للتصدير.

بيستخدم المفتاح الفعلي، فليه تكلفة حقيقية (بضعة سنتات للعقد النموذجي).
"""
from __future__ import annotations

import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import httpx  # noqa: E402

BASE = "http://127.0.0.1:8756/api"
SAMPLE = Path("storage/samples/contract_ar.docx")

TERMS = [
    ("الطرف الأول", "First Party"),
    ("الطرف الثاني", "Second Party"),
    ("عقد تقديم خدمات استشارية", "Consultancy Services Agreement"),
]


def wait_for_job(client: httpx.Client, job_id: str, timeout: int = 300) -> dict:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        job = client.get(f"{BASE}/jobs/{job_id}").json()
        if job["status"] in ("done", "failed"):
            return job
        time.sleep(2)
    raise TimeoutError("المهمة اتأخرت")


def main() -> int:
    failures: list[str] = []

    with httpx.Client(timeout=120.0) as client:
        health = client.get(f"{BASE}/health").json()
        if not health["has_api_key"]:
            print("!! المفتاح مش محمّل في الخادم — أعد تشغيله بعد إنشاء .env")
            return 1

        # ---------- المصطلحات ----------
        existing = {t["source_term"] for t in client.get(f"{BASE}/glossary").json()}
        added = 0
        for source, target in TERMS:
            if source in existing:
                continue
            client.post(
                f"{BASE}/glossary",
                json={
                    "source_term": source,
                    "target_term": target,
                    "domain": "legal",
                    "project_id": None,
                    "notes": "",
                },
            )
            added += 1
        print(f"=== المصطلحات: {added} جديد، {len(existing)} موجود ===")

        # ---------- المشروع ----------
        project = client.post(
            f"{BASE}/projects",
            json={
                "name": "ترجمة حقيقية — عقد استشارات",
                "domain": "legal",
                "model": "claude-sonnet-5",
                "style_notes": "احتفظ بأسماء الشركات كما هي دون ترجمة.",
            },
        ).json()
        print(f"\n=== المشروع: {project['name']} ===")

        # اختبار ترميز: الاسم العربي لازم يرجع زي ما هو
        if project["name"] != "ترجمة حقيقية — عقد استشارات":
            failures.append(f"الاسم العربي اتشوّه: {project['name']!r}")
        else:
            print("  ✓ الترميز العربي سليم في الـ API")

        # ---------- الرفع والاستخراج ----------
        with SAMPLE.open("rb") as handle:
            file = client.post(
                f"{BASE}/projects/{project['id']}/files",
                files={"file": (SAMPLE.name, handle)},
            ).json()
        print(f"\n=== الملف: {file['original_filename']} ===")

        job = client.post(f"{BASE}/files/{file['id']}/extract").json()
        job = wait_for_job(client, job["id"])
        if job["status"] != "done":
            print(f"  ✗ فشل الاستخراج: {job['error']}")
            return 1
        print(f"  استخراج: {job['result']}")

        # ---------- التقدير قبل الصرف ----------
        estimate = client.get(f"{BASE}/projects/{project['id']}/estimate").json()
        print(f"\n=== التقدير المسبق ===")
        print(f"  {estimate['words']} كلمة · {estimate['segments']} مقطع · "
              f"تغطية الذاكرة {estimate['memory_coverage_pct']}%")
        for option in estimate["options"]:
            print(f"    {option['label']:18} ${option['cost_usd']:.4f}")

        # ---------- الترجمة الحقيقية ----------
        print(f"\n=== الترجمة الحقيقية (Sonnet 5) ===")
        start = time.monotonic()
        job = client.post(
            f"{BASE}/files/{file['id']}/translate",
            json={"engine": "claude", "use_memory": True},
        ).json()
        job = wait_for_job(client, job["id"])
        elapsed = time.monotonic() - start

        if job["status"] != "done":
            print(f"  ✗ فشلت الترجمة: {job['error']}")
            return 1

        result = job["result"]
        print(f"  مترجَم={result['translated']} من_الذاكرة={result['from_memory']} "
              f"فشل={result['failed']} تنبيهات={result['flagged']}")
        print(f"  نداءات API={result['api_calls']} · التكلفة=${result['cost_usd']:.6f} "
              f"· الزمن={elapsed:.1f} ثانية")

        if result["failed"]:
            failures.append(f"مقاطع فشلت: {result['failed']}")

        # ---------- عيّنة من المخرجات ----------
        page = client.get(
            f"{BASE}/files/{file['id']}/segments", params={"limit": 300}
        ).json()
        print(f"\n=== عيّنة من الترجمة ===")
        shown = 0
        for segment in page["items"]:
            if not segment["is_translatable"] or not segment["target_text"].strip():
                continue
            print(f"\n  [{segment['location']}]")
            print(f"    ع: {segment['source_text'][:88]}")
            print(f"    E: {segment['target_text'][:88]}")
            if segment["qa_flags"]:
                print(f"    ⚠ {segment['qa_flags']}")
            shown += 1
            if shown >= 7:
                break

        untranslated = [
            s for s in page["items"]
            if s["is_translatable"] and not s["target_text"].strip()
        ]
        if untranslated:
            failures.append(f"مقاطع فاضية بعد الترجمة: {len(untranslated)}")

        flagged = [s for s in page["items"] if s["qa_flags"]]
        print(f"\n=== فحوصات الجودة: {len(flagged)} مقطع عليه تنبيه ===")
        for segment in flagged[:5]:
            print(f"  {segment['location']}: {segment['qa_flags']}")

        # ---------- التصدير ----------
        export = client.post(f"{BASE}/files/{file['id']}/export").json()
        print(f"\n=== التصدير ===")
        print(f"  {export['filename']} ({export['size_bytes']:,} بايت)")

        output = Path("storage/projects") / project["id"] / "output" / export["filename"]
        if not output.exists():
            failures.append("ملف التصدير مش موجود على القرص")
        else:
            from docx import Document

            doc = Document(str(output))
            text = "\n".join(p.text for p in doc.paragraphs)
            text += "\n" + "\n".join(
                c.text for t in doc.tables for r in t.rows for c in r.cells
            )
            arabic = sum(1 for ch in text if "؀" <= ch <= "ۿ")
            latin = sum(1 for ch in text if "a" <= ch.lower() <= "z")
            print(f"  حروف لاتينية={latin} · حروف عربية متبقية={arabic}")

            if latin < 200:
                failures.append("الملف المصدَّر مافيهوش نص إنجليزي كافٍ")
            if arabic > 20:
                failures.append(f"لسه فيه {arabic} حرف عربي في المخرجات")

            print("\n  --- أول أسطر الملف المترجم ---")
            for line in [p.text for p in doc.paragraphs if p.text.strip()][:6]:
                print(f"    {line[:90]}")

        # ---------- تقرير التكلفة ----------
        cost = client.get(f"{BASE}/projects/{project['id']}/cost").json()
        print(f"\n=== تقرير التكلفة الفعلي ===")
        print(f"  الإجمالي: ${cost['total_cost_usd']:.6f}")
        print(f"  لكل كلمة: ${cost['cost_per_word']:.6f} · "
              f"لكل صفحة: ${cost['cost_per_page']:.4f}")
        print(f"  وفورات: {cost['savings']['segments_from_memory']} مقطع من الذاكرة · "
              f"${cost['savings']['cache_saving_usd']:.4f} من التخزين المؤقت")

        for row in cost["by_model"]:
            print(f"    {row['model']}: {row['calls']} نداء · "
                  f"إدخال={row['input_tokens']:,} إخراج={row['output_tokens']:,} "
                  f"كاش_قراءة={row['cache_read_tokens']:,}")

        if cost["total_cost_usd"] <= 0:
            failures.append("التكلفة الفعلية صفر — السجلّ مش شغّال")

        print(f"\n  رابط المشروع: http://127.0.0.1:8756/translator/{project['id']}")

    print("\n" + "=" * 64)
    if failures:
        print(f"فشل: {len(failures)} مشكلة")
        for item in failures:
            print(f"  ✗ {item}")
        return 1
    print("نجح: الترجمة الحقيقية من الرفع للتصدير سليمة ✓")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
