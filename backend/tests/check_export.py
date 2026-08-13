"""التحقق من تصدير أحدث ملف مترجم — الصيغة والتنسيق والاتجاه."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from sqlalchemy import select  # noqa: E402

from app.core.db import SessionLocal  # noqa: E402
from app.models import Project, SourceFile  # noqa: E402
from app.tools.translator import pipeline  # noqa: E402


def main() -> int:
    failures: list[str] = []
    db = SessionLocal()
    try:
        project = db.execute(
            select(Project).order_by(Project.created_at.desc()).limit(1)
        ).scalar_one()
        file = db.execute(
            select(SourceFile).where(SourceFile.project_id == project.id).limit(1)
        ).scalar_one()

        print(f"المشروع: {project.name}")
        print(f"الملف  : {file.original_filename}")

        output = pipeline.export_file(db, file)
        db.commit()
        print(f"\nالمخرَج: {output.name} ({output.stat().st_size:,} بايت)")

        doc = Document(str(output))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        cells = [
            c.text
            for t in doc.tables
            for r in t.rows
            for c in r.cells
            if c.text.strip()
        ]
        combined = "\n".join(paragraphs + cells)

        arabic = sum(1 for ch in combined if "؀" <= ch <= "ۿ")
        latin = sum(1 for ch in combined if "a" <= ch.lower() <= "z")
        print(f"فقرات={len(paragraphs)} خلايا={len(cells)}")
        print(f"حروف لاتينية={latin:,} · حروف عربية متبقية={arabic}")

        if latin < 500:
            failures.append("النص الإنجليزي قليل بشكل مريب")
        if arabic > 30:
            failures.append(f"لسه فيه {arabic} حرف عربي في المتن")

        # الاتجاه اتقلب؟
        bidi = len(list(doc.element.iter(qn("w:bidi"))))
        rtl = len(list(doc.element.iter(qn("w:rtl"))))
        print(f"عناصر اتجاه عربي متبقية: bidi={bidi} rtl={rtl}")
        if bidi or rtl:
            failures.append(f"قلب الاتجاه ناقص: bidi={bidi} rtl={rtl}")

        # الجداول محفوظة؟
        print(f"جداول: {len(doc.tables)}")
        if doc.tables:
            first = doc.tables[0]
            print(f"  أول جدول: {len(first.rows)} صف × {len(first.columns)} عمود")

        print("\n--- أول ١٢ سطر من الملف المترجم ---")
        for line in paragraphs[:12]:
            print(f"  {line[:95]}")

        print("\n--- عيّنة من الجداول ---")
        for line in cells[:10]:
            print(f"  {line[:95]}")

    finally:
        db.close()

    print("\n" + "=" * 62)
    if failures:
        print(f"فشل: {len(failures)}")
        for item in failures:
            print(f"  ✗ {item}")
        return 1
    print("نجح: التصدير سليم ✓")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
